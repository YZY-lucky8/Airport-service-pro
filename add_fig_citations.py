#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""给5个缺失引用的图添加正文描述，让图自然融入正文"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

def insert_paragraph_before(para, text):
    """在段落前插入新段落"""
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, doc)
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return new_para

# ============================================================
# 图2-1 系统架构总览图（段落97）
# 在2.1节标题后添加介绍段落
# ============================================================
print("=== 图2-1 系统架构总览图 ===")
# 段落95是"2.1 系统总体设计方案"，在其后插入介绍
para_21 = doc.paragraphs[95]
insert_paragraph_before(doc.paragraphs[97], 
    "本章详细阐述系统的设计方案与实现细节。系统整体架构如图2-1所示，采用分层设计思想，将安全防护与智能服务深度融合，确保各模块职责清晰、耦合度低。")
print("  ✅ 已在图2-1前添加正文引用")

# ============================================================
# 图2-2 系统总体架构与技术路线（段落113，插入后索引会变，需要重新定位）
# ============================================================
print("\n=== 图2-2 系统总体架构与技术路线 ===")
# 重新找到图2-2
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('图 2-2'):
        # 在段落108（三级架构描述）中添加引用
        for j in range(max(0, i-10), i):
            text = doc.paragraphs[j].text
            if '三级轻量化架构' in text and '如图' not in text:
                set_para_text(doc.paragraphs[j], 
                    text.replace('三级轻量化架构，层次清晰、分工明确。', 
                                  '三级轻量化架构，层次清晰、分工明确，系统总体架构与技术路线如图2-2所示。'))
                print(f"  ✅ 段落{j}: 已添加图2-2引用")
                break
        break

# ============================================================
# 图2-3 机场安全运营中心态势总览页面
# ============================================================
print("\n=== 图2-3 机场安全运营中心态势总览页面 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('图 2-3'):
        insert_paragraph_before(para,
            "安全防护模块是系统的核心组成部分，机场安全运营中心态势总览页面如图2-3所示，实时展示系统运行状态、攻击告警、防护效果等关键信息。")
        print("  ✅ 已在图2-3前添加正文引用")
        break

# ============================================================
# 图2-10 主界面地图导航页
# ============================================================
print("\n=== 图2-10 主界面地图导航页 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('图 2-10'):
        insert_paragraph_before(para,
            "核心功能模块为旅客提供智能化出行服务，主界面地图导航页如图2-10所示，集成了航班查询、地点导航、天气查询、知识库问答等功能入口。")
        print("  ✅ 已在图2-10前添加正文引用")
        break

# ============================================================
# 图2-11 旅客端自助选座界面
# ============================================================
print("\n=== 图2-11 旅客端自助选座界面 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('图 2-11'):
        # 在前一段值机描述中添加引用
        for j in range(max(0, i-5), i):
            text = doc.paragraphs[j].text
            if '可视化机舱布局选座界面' in text and '如图' not in text:
                set_para_text(doc.paragraphs[j],
                    text.replace('加载可视化机舱布局选座界面',
                                  '加载可视化机舱布局选座界面，旅客端自助选座界面如图2-11所示'))
                print(f"  ✅ 段落{j}: 已添加图2-11引用")
                break
        break

# ============================================================
# 验证所有图都有引用
# ============================================================
print("\n=== 验证所有图引用 ===")
import re
fig_captions = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^图\s+\d+-\d+', text):
        fig_captions.append(text)

fig_cited = set()
for para in doc.paragraphs:
    text = para.text
    if not re.match(r'^图\s+\d+-\d+', text.strip()):
        for m in re.findall(r'图\s*(\d+-\d+)', text):
            fig_cited.add(m)

all_cited = True
for cap in fig_captions:
    fid = re.match(r'^图\s+(\d+-\d+)', cap).group(1)
    status = "✅" if fid in fig_cited else "❌"
    if fid not in fig_cited:
        all_cited = False
    print(f"  {status} {cap[:40]}")

if all_cited:
    print("\n✅ 所有14个图都有正文引用！")
else:
    print("\n⚠️ 仍有图缺少引用")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')
print("\n✅ 已保存")
