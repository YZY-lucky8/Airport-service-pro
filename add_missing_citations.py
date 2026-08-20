#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""补充缺失的4个图表引用"""
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

# 查看缺失引用的图表前后段落
targets = [
    (188, 'fig', '2-10', '图 2-10 主界面地图导航页'),
    (229, 'tbl', '2-2', '表 2-2 智能体扩展数据表清单'),
    (430, 'fig', '3-2', '图 3-2 ipwhitelist.js测试结果'),
    (432, 'fig', '3-3', '图 3-3 unifyResponse.js'),
]

print("=== 查看缺失引用图表的上下文 ===")
for para_idx, ftype, fid, caption in targets:
    print(f"\n--- {caption} (段落{para_idx}) ---")
    # 前面5段
    for j in range(max(0, para_idx-5), para_idx):
        text = doc.paragraphs[j].text.strip()
        if text:
            print(f"  前[{j}]: {text[:80]}")
    # 后面2段
    for j in range(para_idx+1, min(len(doc.paragraphs), para_idx+3)):
        text = doc.paragraphs[j].text.strip()
        if text:
            print(f"  后[{j}]: {text[:80]}")

# 补充引用
print("\n=== 补充引用 ===")

# 图2-10 主界面地图导航页 - 在2.4.2自然语言地点导航的描述中添加
# 表2-2 智能体扩展数据表清单 - 在数据库设计描述中添加
# 图3-2 ipwhitelist.js测试结果 - 在3.6.1 DDoS防护测试中添加
# 图3-3 unifyResponse.js - 在3.6.1 DDoS防护测试中添加

# 遍历文档流找到合适位置
body = doc.element.body
children = list(body.iterchildren())

# 定义需要补充的引用：(图题段落索引, 引用类型, 编号, 要添加引用的段落关键词)
supplements = [
    # 图2-10: 找地点导航相关描述
    ('fig', '2-10', '地点导航', '自然语言地点导航'),
    # 表2-2: 找扩展表相关描述
    ('tbl', '2-2', '扩展表', '智能体扩展'),
    # 图3-2: 找ipwhitelist测试描述
    ('fig', '3-2', 'ipwhitelist', '白名单'),
    # 图3-3: 找unifyResponse测试描述
    ('fig', '3-3', 'unifyResponse', '统一响应'),
]

added = 0
for ftype, fid, keyword1, keyword2 in supplements:
    citation_text = f"如图{fid}所示" if ftype == 'fig' else f"如表{fid}所示"
    
    # 找包含关键词的段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if (keyword1 in text or keyword2 in text) and citation_text not in text and not re.match(r'^图\s', text) and not re.match(r'^表\s', text):
            # 检查是否已经有同类型引用
            if ftype == 'fig' and '如图' in text:
                continue
            if ftype == 'tbl' and '如表' in text:
                continue
            # 添加引用
            new_text = text + f"（{citation_text}）"
            set_para_text(para, new_text)
            print(f"  ✅ {citation_text}: 段落{i} - {text[:50]}...")
            added += 1
            break

print(f"\n共补充 {added} 处引用")

# 再次验证
print("\n=== 最终验证 ===")
all_citations = set()
for para in doc.paragraphs:
    text = para.text
    for m in re.findall(r'如图\s*(\d+-\d+)', text):
        all_citations.add(('fig', m))
    for m in re.findall(r'如表\s*(\d+-\d+)', text):
        all_citations.add(('tbl', m))

all_figs = set()
all_tables = set()
for para in doc.paragraphs:
    text = para.text.strip()
    m = re.match(r'^图\s+(\d+-\d+)', text)
    if m: all_figs.add(m.group(1))
    m = re.match(r'^表\s+(\d+-\d+)', text)
    if m: all_tables.add(m.group(1))

missing_figs = all_figs - set(f for t, f in all_citations if t == 'fig')
missing_tables = all_tables - set(t for t, f in all_citations if t == 'tbl')

print(f"图: {len(all_figs)}个, 已引用: {len([f for t,f in all_citations if t=='fig'])}个, 缺失: {missing_figs}")
print(f"表: {len(all_tables)}个, 已引用: {len([f for t,f in all_citations if t=='tbl'])}个, 缺失: {missing_tables}")

if not missing_figs and not missing_tables:
    print("✅ 所有图表都有正文引用！")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')
print("\n✅ 已保存")
