#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""详细列出所有图表引用的上下文"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3).docx')

print("=== 所有括号模式引用 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    matches = re.findall(r'（如[图表]\d+-\d+所示）', text)
    if matches:
        print(f"\n[{i}] {text[:150]}")
        print(f"    引用: {matches}")

print("\n\n=== 表2-2内容 ===")
# 找到表2-2
for ti, table in enumerate(doc.tables):
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if '扩展表名称' in str(first_row) or '扩展表' in str(first_row):
        print(f"表{ti}:")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip()[:30] for cell in row.cells]
            print(f"  行{ri}: {cells}")
        break

print("\n\n=== 3.2.1节内容 ===")
in_321 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.2.1' in text:
        in_321 = True
    if in_321:
        print(f"[{i}] {text[:100]}")
    if in_321 and '3.2.2' in text:
        break

print("\n\n=== 表3-7前后内容 ===")
for i, para in enumerate(doc.paragraphs):
    if 369 <= i <= 375:
        print(f"[{i}] {para.text[:120]}")
