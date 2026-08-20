#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析文档结构，定位需要修改的段落和表格"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告.docx')

print("=== 段落分析 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(kw in text for kw in ['2.7.4', '性能指标', '部署效率', '部署时间', '第四章', '4.1', '理念创新', '安全与服务一体化', '国产化自主可控', '适老化普惠']):
        print(f"[{i}] style={para.style.name} | {text[:100]}")

print("\n=== 表格分析 ===")
for ti, table in enumerate(doc.tables):
    # 看第一行内容判断是什么表
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    print(f"\n表{ti}: 行数={len(table.rows)}, 列数={len(table.columns)}")
    print(f"  首行: {first_row}")
    # 如果是性能指标表，打印所有内容
    if any('指标' in c or '并发' in c or '响应时间' in c for c in first_row):
        print("  >>> 这是性能指标表！")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  行{ri}: {cells}")
