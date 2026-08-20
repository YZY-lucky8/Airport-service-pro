#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析文档中所有图、表的位置和当前编号"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\作品报告-08.15-2.docx')

print("=== 文档章节结构 ===")
chapters = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    if style.startswith('Heading 1') and text:
        chapters.append((i, text))
        print(f"[{i}] {text}")

print(f"\n=== 所有图题（段落中含'图'+数字）===")
figures = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # 匹配图题：图X 或 图 X 开头
    if re.match(r'^图\s*\d+', text):
        figures.append((i, text[:100]))
        # 确定属于哪一章
        chap = "未知"
        for ci, ct in chapters:
            if ci < i:
                chap = ct
        print(f"[{i}] [{chap[:15]}] {text[:100]}")

print(f"\n=== 所有表题（段落中含'表'+数字）===")
tables_caption = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^表\s*\d+', text):
        tables_caption.append((i, text[:100]))
        chap = "未知"
        for ci, ct in chapters:
            if ci < i:
                chap = ct
        print(f"[{i}] [{chap[:15]}] {text[:100]}")

print(f"\n=== 表格总数: {len(doc.tables)} ===")
# 检查每个表格前后的表题
for ti, table in enumerate(doc.tables):
    # 找表格前最近的表题段落
    # 由于python-docx无法直接获取表格在文档流中的位置，我们通过段落和表格的混合遍历来判断
    pass

# 用另一种方式：遍历文档body中的所有元素
print("\n=== 文档流中图表顺序 ===")
from docx.oxml.ns import qn
body = doc.element.body
fig_count = 0
table_count = 0
current_chap = "前言"
item_idx = 0

for child in body.iterchildren():
    if child.tag == qn('w:p'):
        # 段落
        from docx.text.paragraph import Paragraph
        para = Paragraph(child, doc)
        text = para.text.strip()
        # 检查是否是章节标题
        for ci, ct in chapters:
            if para._element is doc.paragraphs[ci]._element:
                current_chap = ct
                break
        # 检查是否是图题
        if re.match(r'^图\s*\d+', text):
            fig_count += 1
            print(f"  图{fig_count}: [{current_chap[:20]}] {text[:80]}")
        # 检查是否是表题
        if re.match(r'^表\s*\d+', text):
            table_count += 1
            print(f"  表{table_count}: [{current_chap[:20]}] {text[:80]}")
    elif child.tag == qn('w:tbl'):
        # 表格
        table_count += 1
        # 找表格前一个段落作为表题
        pass

print(f"\n总计: 图{fig_count}个, 表{table_count}个（含表格对象）")
