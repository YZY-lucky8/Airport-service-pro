#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析新文件结构：图表、引用、占位符"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\作品报告-(3) - 用于合并.docx')

print("=== 章节结构 ===")
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading') and para.text.strip():
        print(f"[{i}] {para.text[:50]}")

print(f"\n=== 表格数: {len(doc.tables)} ===")

print("\n=== 所有图题表题 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^[图表]\s*\d', text):
        print(f"[{i}] {text[:60]}")

print("\n=== 括号模式引用 ===")
count = 0
for i, para in enumerate(doc.paragraphs):
    if re.search(r'（如[图表]\d', para.text):
        count += 1
        print(f"[{i}] {para.text[:80]}")
print(f"共{count}处")

print("\n=== 小圆圈占位符 ===")
circle = 0
for i, para in enumerate(doc.paragraphs):
    if '⃝' in para.text or '○' in para.text:
        circle += 1
        print(f"[{i}] {para.text[:60]}")
print(f"共{circle}处")

print("\n=== 图3-1到3-3引用检查 ===")
for i, para in enumerate(doc.paragraphs):
    if '图3' in para.text or '图 3' in para.text:
        print(f"[{i}] {para.text[:80]}")
