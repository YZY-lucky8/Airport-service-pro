#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析海报要求文档内容和图片"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\附件3-路演海报要求(4).docx')

print("=== 文档段落内容 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i}] ({para.style.name}) {text[:300]}")

print(f"\n=== 文档表格 ===")
print(f"表格数量: {len(doc.tables)}")
for ti, table in enumerate(doc.tables):
    print(f"\n表{ti}: {len(table.rows)}行 x {len(table.columns)}列")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:80] for cell in row.cells]
        print(f"  行{ri}: {cells}")

print(f"\n=== 文档图片 ===")
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
        print(f"  图片{img_count}: {rel.target_ref}")

if img_count == 0:
    print("  文档中没有嵌入图片")
