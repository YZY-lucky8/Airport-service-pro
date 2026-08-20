#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析报告结构：参考文献 + 正文引用位置 + 潜在问题"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告.docx')

# 1. 找到参考文献章节
print("=== 参考文献章节 ===")
in_ref = False
ref_start = -1
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '参考文献' in text and len(text) < 20:
        in_ref = True
        ref_start = i
        print(f"\n[{i}] {text} (样式: {para.style.name})")
        continue
    if in_ref:
        if text:
            print(f"[{i}] {text[:150]}")
        # 遇到下一个大标题停止
        if para.style.name.startswith('Heading') and i > ref_start + 2:
            break

# 2. 统计正文中已有角标引用
print("\n\n=== 正文中已有角标引用 ===")
import re
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # 匹配 [数字] 形式的角标
    matches = re.findall(r'\[(\d+)\]', text)
    if matches:
        print(f"[{i}] 引用{matches}: {text[:100]}")

# 3. 检查可能需要引用的技术描述段落
print("\n\n=== 可能需要引用参考文献的技术段落 ===")
keywords = ['零信任', 'DDoS', '布隆过滤器', 'Bloom', '智能体', 'Agent', 'RAG', 
            'HMAC', 'JWT', '滑动窗口', 'Prompt注入', '大语言模型', 'LLM',
            '深度学习', '机器学习', '区块链', '边缘计算', '物联网', '5G',
            'Web Speech', '语音识别', '情感分析', '知识库']
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text or len(text) < 30:
        continue
    for kw in keywords:
        if kw in text and '[' not in text[:5]:
            # 检查是否已有角标
            if not re.search(r'\[\d+\]', text):
                print(f"[{i}] 含'{kw}': {text[:80]}...")
                break

# 4. 检查表格数量和关键表格
print(f"\n\n=== 文档统计 ===")
print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")
