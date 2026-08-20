#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析作品报告-(3).docx的结构"""
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3).docx')

print("=== 章节结构 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    if style.startswith('Heading') and text:
        print(f"[{i}] ({style}) {text[:60]}")

print(f"\n=== 表格数量: {len(doc.tables)} ===")

print("\n=== 所有图题和表题 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^图\s*\d', text) or re.match(r'^表\s*\d', text):
        print(f"[{i}] {text[:80]}")

print("\n=== 检查小圆圈占位符 ===")
circle_count = 0
for i, para in enumerate(doc.paragraphs):
    if '⃝' in para.text or '○' in para.text:
        circle_count += 1
        print(f"[{i}] {para.text[:80]}")
print(f"共找到 {circle_count} 个含占位符的段落")

print("\n=== 检查1.2.2节的备注 ===")
in_122 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1.2.2' in text:
        in_122 = True
    if in_122 and ('注' in text or '备注' in text):
        print(f"[{i}] {text[:120]}")
    if in_122 and text.startswith('1.3'):
        break

print("\n=== 检查图表引用方式 ===")
bracket_cite = 0
natural_cite = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if re.search(r'（如[图表]\d', text):
        bracket_cite += 1
    if re.search(r'[，。][如图表]\d', text) or re.search(r'^[如图表]\d', text):
        natural_cite += 1
print(f"括号模式引用: {bracket_cite} 处")
print(f"自然引入引用: {natural_cite} 处")
