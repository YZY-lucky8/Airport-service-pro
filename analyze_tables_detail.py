#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""详细分析每个表格的位置和前后文"""
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

doc = Document(r'C:\Users\Lenovo\Desktop\作品报告-08.15-2.docx')

chapters = []
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading 1') and para.text.strip():
        chapters.append((i, para.text.strip()))

# 先构建文档流列表
print("=== 文档流中所有表格的位置和上下文 ===")
body = doc.element.body
current_chap = "前言"
table_idx = 0

# 收集所有子元素
children = list(body.iterchildren())

for idx, child in enumerate(children):
    if child.tag == qn('w:p'):
        para = Paragraph(child, doc)
        text = para.text.strip()
        for ci, ct in chapters:
            if para._element is doc.paragraphs[ci]._element:
                current_chap = ct
                break
    elif child.tag == qn('w:tbl'):
        table_idx += 1
        tbl = Table(child, doc)
        first_row = [cell.text.strip()[:20] for cell in tbl.rows[0].cells] if tbl.rows else []
        rows_count = len(tbl.rows)
        cols_count = len(tbl.columns)
        
        # 找前面的非空段落
        prev_text = ""
        for j in range(idx-1, max(0, idx-5), -1):
            if children[j].tag == qn('w:p'):
                p = Paragraph(children[j], doc)
                if p.text.strip():
                    prev_text = p.text.strip()[:60]
                    break
        
        # 找后面的非空段落
        next_text = ""
        for j in range(idx+1, min(len(children), idx+5)):
            if children[j].tag == qn('w:p'):
                p = Paragraph(children[j], doc)
                if p.text.strip():
                    next_text = p.text.strip()[:60]
                    break
        
        print(f"\n表{table_idx}: [{current_chap[:15]}] {rows_count}行x{cols_count}列")
        print(f"  表头: {first_row}")
        print(f"  前文: {prev_text}")
        print(f"  后文: {next_text}")

print(f"\n总计: {table_idx}个表格")
