#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""检查所有图的引用情况和图片位置"""
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

print("=== 所有图题及前后文 ===")
fig_captions = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^图\s+\d+-\d+', text):
        fig_captions.append((i, text))

# 收集正文中的图引用
fig_cited = set()
for para in doc.paragraphs:
    for m in re.findall(r'图\s*(\d+-\d+)', para.text):
        if not re.match(r'^图\s+\d+-\d+', para.text.strip()):
            fig_cited.add(m)

for para_idx, caption in fig_captions:
    fig_id = re.match(r'^图\s+(\d+-\d+)', caption).group(1)
    has_citation = fig_id in fig_cited
    
    # 找前面的引用段落
    cite_para = "无"
    for j in range(max(0, para_idx-10), para_idx):
        text = doc.paragraphs[j].text
        if f'图{fig_id}' in text or f'图 {fig_id}' in text:
            cite_para = f"段落{j}: {text[:50]}..."
            break
    
    # 找图片所在段落（图题前一段通常是图片）
    has_image = False
    if para_idx > 0:
        prev_para = doc.paragraphs[para_idx - 1]
        # 检查段落中是否有图片
        if 'graphic' in prev_para._element.xml or 'pic:pic' in prev_para._element.xml:
            has_image = True
    
    status = "✅" if has_citation else "❌"
    print(f"\n{status} {caption}")
    print(f"   图题段落: {para_idx}")
    print(f"   有正文引用: {has_citation}")
    print(f"   引用位置: {cite_para}")
    print(f"   图题前一段有图片: {has_image}")

print(f"\n=== 统计 ===")
print(f"图总数: {len(fig_captions)}")
print(f"已引用: {len(fig_cited)}")
missing = [fid for _, cap in fig_captions for fid in [re.match(r'^图\s+(\d+-\d+)', cap).group(1)] if fid not in fig_cited]
print(f"未引用: {missing}")
