#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""检查报告中已有的图号引用"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_最终版.docx')

print("=== 报告中已有的图号引用 ===")
fig_nums = set()
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # 匹配"图X"或"图 X"格式
    matches = re.findall(r'图\s*(\d+)', text)
    for m in matches:
        fig_nums.add(int(m))
        if len(text) < 100:
            print(f"  段落{i}: 图{m} - {text.strip()[:80]}")

print(f"\n已使用的图号: {sorted(fig_nums)}")
print(f"最大图号: {max(fig_nums) if fig_nums else '无'}")

# 检查2.1.3、2.2.1、2.2.3、2.4节的具体段落
print("\n=== 关键章节段落内容 ===")
key_sections = {
    '2.1.3 整体架构': 116,
    '2.2.1 六项DDoS防护机制末尾': 143,
    '2.2.3 三道智能体安全防线末尾': 152,
    '2.4 核心功能模块实现': 178,
}
for name, idx in key_sections.items():
    if idx < len(doc.paragraphs):
        text = doc.paragraphs[idx].text.strip()
        print(f"\n[{idx}] {name}:")
        print(f"  {text[:150]}")
        # 看下一段
        if idx + 1 < len(doc.paragraphs):
            print(f"  下一段[{idx+1}]: {doc.paragraphs[idx+1].text.strip()[:100]}")
