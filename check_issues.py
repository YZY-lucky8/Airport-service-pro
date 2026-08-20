#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""检查报告中与代码可能不一致的关键描述"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告.docx')

print("=== 1. 布隆过滤器描述检查 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '布隆过滤器' in text and ('哈希' in text or '位' in text or '容量' in text):
        print(f"[{i}] {text[:200]}")

print("\n=== 2. 第五章测试结果检查 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if ('功能测试' in text or '通过率' in text or '测试用例' in text) and ('33' in text or '26' in text or '78' in text or '36' in text):
        print(f"[{i}] {text[:200]}")

print("\n=== 3. 智能体代码行数描述检查 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if ('行' in text and ('代码' in text or 'dispatcher' in text or '智能体' in text)) and re.search(r'\d+\s*行', text):
        print(f"[{i}] {text[:200]}")

print("\n=== 4. 性能指标关键数据检查 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if any(x in text for x in ['70MB', '70兆', '12KB', '95.3%', '93%', '163/171']):
        print(f"[{i}] {text[:200]}")

print("\n=== 5. 表格中的性能数据 ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        row_text = ' | '.join(cells)
        if any(x in row_text for x in ['70MB', '12KB', '95.3%', '93%', '163', '70兆']):
            print(f"表{ti}行{ri}: {row_text[:200]}")

print("\n=== 6. 第五章5.2节内容 ===")
in_52 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '5.2' in text and ('测试' in text or '结果' in text):
        in_52 = True
    if in_52 and text:
        print(f"[{i}] {text[:150]}")
    if in_52 and ('5.3' in text or '第六章' in text):
        break
