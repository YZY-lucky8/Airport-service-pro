#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""查看缺失引用的5个图的上下文"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

targets = [97, 113, 119, 168, 189]  # 图2-1, 2-2, 2-3, 2-10, 2-11

for para_idx in targets:
    print(f"\n=== 图题段落{para_idx}: {doc.paragraphs[para_idx].text[:50]} ===")
    # 前面5段
    for j in range(max(0, para_idx-5), para_idx):
        text = doc.paragraphs[j].text.strip()
        if text:
            print(f"  前[{j}]: {text[:80]}")
    # 后面3段
    for j in range(para_idx+1, min(len(doc.paragraphs), para_idx+4)):
        text = doc.paragraphs[j].text.strip()
        if text:
            print(f"  后[{j}]: {text[:80]}")
