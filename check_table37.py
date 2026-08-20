#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""查看表3-7的实际内容"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx')

# 找到表3-7（索引12）
table = doc.tables[12]
print("=== 表3-7 三级混合检索引擎测试结果表 ===")
for ri, row in enumerate(table.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{ri}: {cells}")

# 查看表3-7前后段落
print("\n=== 表3-7前后段落 ===")
for i, para in enumerate(doc.paragraphs):
    if '3.4.3' in para.text or '三级检索' in para.text or '检索引擎工作稳定' in para.text:
        print(f"[{i}] {para.text[:150]}")
