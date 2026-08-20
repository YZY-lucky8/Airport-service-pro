#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
综合修改脚本：处理作品报告-(3).docx的19项问题
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re
import copy

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3).docx')

def set_para_text(para, text):
    """设置段落文本，保留第一个run格式"""
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

def insert_paragraph_after(para, text, style=None):
    """在段落后插入新段落"""
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    new_para = Paragraph(new_p, doc)
    if text:
        run = new_para.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if style:
        new_para.style = doc.styles[style]
    return new_para

def set_table_header_repeat(table):
    """设置表格第一行为重复标题行（跨页时自动重复）"""
    tr = table.rows[0]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

# ============================================================
# 1. 优化图表引用方式（括号模式 → 自然引入）
# ============================================================
print("=== 1. 优化图表引用方式 ===")
citation_fixes = [
    (44, "系统各项核心指标均达到或超过设计目标：（如表1-1所示）", "系统各项核心指标均达到或超过设计目标，如表1-1所示："),
    (81, "本项目目标客户聚焦于公共服务场景运营方与渠道合作方：（如表1-2所示）", "本项目目标客户聚焦于公共服务场景运营方与渠道合作方，如表1-2所示："),
    (124, "（一）滑动窗口频率检测。（如图2-4所示）该机制采用", "（一）滑动窗口频率检测。如图2-4所示，该机制采用"),
    (127, "布隆过滤器IP黑名单。（如图2-5所示）被滑动窗口检测拦截的IP地址自动加入黑名单", "布隆过滤器IP黑名单。图2-5展示被滑动窗口检测拦截的IP地址自动加入黑名单的机制，后续请求通过布隆过滤器进行微秒级查询拦截"),
    (131, "（三）HMAC一次性令牌。（如图2-6所示）该机制基于", "（三）HMAC一次性令牌。如图2-6所示，该机制基于"),
    (136, "（六）关键服务免扰。（如图2-7所示） 紧急帮助接口", "（六）关键服务免扰。如图2-7所示，紧急帮助接口"),
    (139, "请求成功率超过99.2%。（如图2-8所示）", "请求成功率超过99.2%，如图2-8所示。"),
    (209, "各表存储用途、核心字段、预置数据量整理如下表：（如表2-1所示）", "各表存储用途、核心字段、预置数据量整理如下，如表2-1所示："),
    (227, "实测，系统核心性能指标如下：（如表2-3所示）", "实测，系统核心性能指标如下，如表2-3所示："),
    (262, "知识库实际包含三十三条条目，覆盖七分类：（如表3-2所示）", "知识库实际包含三十三条条目，覆盖七分类，如表3-2所示："),
    (281, "航班查询测试小结：（如表3-3所示）", "航班查询测试小结如表3-3所示："),
    (308, "知识库问答测试小结：（如表3-4所示）", "知识库问答测试小结如表3-4所示："),
    (363, "系统定义了九类预定义意图，经实际测试验证：（如表3-5所示）", "系统定义了九类预定义意图，经实际测试验证，结果如表3-5所示："),
    (370, "双字符分词处理中文文本。（如表3-7所示）", "双字符分词处理中文文本，测试结果如表3-7所示："),
    (419, "通过接口响应 latency 字段实测：（如表3-9所示）", "通过接口响应 latency 字段实测，结果如表3-9所示："),
    (478, "对比优势：（如表4-1所示）", "本系统所提方案与传统商业WAF、单机制轻量方案的对比如表4-1所示："),
    (485, "对比优势：（如表4-2所示）", "本系统三道防线方案与传统输入过滤方案、VeriGuard形式化验证方案的对比如表4-2所示："),
    (492, "对比优势：（如表4-3所示）", "本系统三级检索方案与向量数据库RAG、简单关键词匹配方案的对比如表4-3所示："),
    (500, "对比优势：（如表4-4所示）", "本系统自研框架与LangChain/LangGraph、CrewAI、AutoGen等主流框架的对比如表4-4所示："),
]

for para_idx, old, new in citation_fixes:
    if para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        if old in para.text:
            set_para_text(para, para.text.replace(old, new))
            print(f"  ✅ 段落{para_idx}: 引用方式已优化")
        else:
            print(f"  ⚠️  段落{para_idx}: 未找到匹配文本")
            print(f"     实际: {para.text[:80]}")

# ============================================================
# 2. 删除1.2.2节的"注："和"备注："
# ============================================================
print("\n=== 2. 删除1.2.2节备注 ===")
# 段落51: 注：以上为核心性能指标概览...
# 段落53: 备注：部署时间依据...
# 将备注内容合并到正文或删除
for i in [51, 53]:
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text.startswith('注：') or text.startswith('备注：'):
            # 删除该段落
            para._element.getparent().remove(para._element)
            print(f"  ✅ 段落{i}: 已删除'{text[:30]}...'")

# ============================================================
# 3. 表2-2第三列添加表头
# ============================================================
print("\n=== 3. 表2-2第三列表头 ===")
# 表2-2是doc.tables[3]
table22 = doc.tables[3]
cell = table22.rows[0].cells[2]
cell.text = ''
p = cell.paragraphs[0]
run = p.add_run('核心字段与说明')
run.font.name = '宋体'
run.font.size = Pt(10)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
print("  ✅ 表2-2第三列表头已添加'核心字段与说明'")

# ============================================================
# 4. 设置跨页表格重复标题行
# ============================================================
print("\n=== 4. 设置跨页表格重复标题行 ===")
# 表2-1(索引2), 表2-3(索引4), 表3-2(索引6), 表3-3(索引7), 表3-15(索引19)等可能跨页
repeat_tables = [2, 4, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 19, 20, 21]
for ti in repeat_tables:
    if ti < len(doc.tables):
        set_table_header_repeat(doc.tables[ti])
        print(f"  ✅ 表{ti}: 已设置重复标题行")

# ============================================================
# 5. 补充只有表格无内容的章节
# ============================================================
print("\n=== 5. 补充表格前导文字 ===")

# 3.2.1 测试环境配置 - 段落257后插入
para_321 = doc.paragraphs[257]
new_para = insert_paragraph_after(para_321, "本节介绍系统测试环境的具体配置，包括服务器硬件、操作系统、运行时环境以及数据库等，确保测试结果可复现，具体配置如表3-1所示。")
print("  ✅ 3.2.1节: 已补充测试环境介绍")

# 3.4.2 情感分析模块测试 - 段落366后插入
para_342 = doc.paragraphs[366]
new_para = insert_paragraph_after(para_342, "本节介绍情感分析模块的测试方法，通过输入带有一定感情色彩的日常对话文本，系统进行情感预判，然后与测试样本的标注标签对比，验证情感识别的准确性，测试结果如表3-6所示。")
print("  ✅ 3.4.2节: 已补充情感分析测试介绍")

# ============================================================
# 6. 表3-7补充说明（为什么只有第一级命中）
# ============================================================
print("\n=== 6. 表3-7补充说明 ===")
# 段落372是表3-7后的"检索引擎工作稳定..."
para_372 = doc.paragraphs[372]
old_text = para_372.text
new_text = old_text + " 表中所有测试用例均在第一级精确匹配命中，原因在于测试用例均使用知识库条目标题中的标准关键词进行查询，符合三级检索引擎的设计预期——标准查询优先在第一级精确匹配，仅当精确匹配无结果时才降级至第二级模糊匹配和第三级全文检索。"
set_para_text(para_372, new_text)
print("  ✅ 表3-7后: 已补充第一级命中原因说明")

# ============================================================
# 7. 补充缺失的图表引用
# ============================================================
print("\n=== 7. 补充缺失图表引用 ===")

# 图2-9 Agent三阶段安全防护 - 在2.2.3节描述中添加
# 段落146是2.2.3标题，找三道防线描述段落
for i in range(146, 153):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text
        if '三道防线' in text and '如图' not in text:
            set_para_text(doc.paragraphs[i], text + "，整体运行机制与安全验证效果如图2-9所示。")
            print(f"  ✅ 段落{i}: 已补充图2-9引用")
            break

# 图3-1到3-3 - 在3.6.1节DDoS防护测试中添加
for i in range(397, 410):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text
        if '白名单' in text and '关键服务免扰' in text and '如图3-1' not in text:
            set_para_text(doc.paragraphs[i], text + " 相关测试结果分别如图3-1、图3-2、图3-3所示。")
            print(f"  ✅ 段落{i}: 已补充图3-1/3-2/3-3引用")
            break

# ============================================================
# 8. 表4-1到4-4前补充介绍（已在引用优化中处理）
# ============================================================
print("\n=== 8. 表4-1到4-4介绍（已在引用优化中处理）===")
print("  ✅ 已在第1步中处理")

# ============================================================
# 9. 每个测试表格后补充结果分析
# ============================================================
print("\n=== 9. 补充测试表格结果分析 ===")

# 需要在表格后添加分析的段落索引（表题段落）
# 表3-3(282), 表3-4(309), 表3-5(364), 表3-6(367), 表3-8(416), 表3-9(420), 表3-10(424), 表3-11(429)
# 表3-15(443), 表3-16(445), 表3-17(447)

# 检查这些表题后是否已有分析文字
table_analysis = {
    282: "由表3-3可知，航班查询功能在各类测试场景下均能正确提取航班号并返回完整信息，标准航班号查询延迟在1ms以内，仅数字开头的特殊航班号因正则覆盖不全需引导用户补充信息，整体功能稳定可靠。",
    309: "由表3-4可知，知识库问答在各类查询场景下均能正确命中相关条目，三级检索引擎工作稳定，响应延迟在1-3ms，满足实时交互需求。",
    364: "由表3-5可知，九类预定义意图在测试场景中全部正确识别，意图分类准确率达到100%，复合意图融合机制工作正常。",
    367: "由表3-6可知，五类情感（焦虑、愤怒、困惑、高兴、恐惧）在测试场景中识别全部正确，高情感场景自动触发安抚策略，情感分析模块性能达标。",
    416: "由表3-8可知，三道智能体安全防线在各类攻击场景下均能正确拦截，输入过滤、运行时监控、输出过滤三层防护协同工作，无漏防误报。",
    420: "由表3-9可知，智能体规则引擎在各类业务意图下处理延迟均低于5ms，平均延迟约1ms，远优于≤10ms的设计目标。",
    424: "由表3-10可知，系统稳定运行内存约77MB，启动内存约51MB，压测峰值约96MB，均在≤100MB的设计目标内。",
    429: "由表3-11可知，老年模式在WCAG 2.1第二级各项条款上均符合要求，字号、按钮尺寸、对比度等指标全面超过标准要求。",
    443: "由表3-15可知，36项功能测试全部通过，通过率100%，覆盖智能体、安全防护、旅客服务、管理端等全部核心模块。",
    445: "由表3-16可知，六项DDoS防护机制和三道智能体安全防线在各类攻击场景下全部成功防御，安全防护能力达标。",
    447: "由表3-17可知，各项性能指标均达到或超过设计目标，系统在并发、延迟、内存、防护损耗等方面表现优异。",
}

# 注意：插入段落会改变索引，所以从后往前插入
for para_idx in sorted(table_analysis.keys(), reverse=True):
    if para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        # 检查下一段是否已有分析
        next_text = doc.paragraphs[para_idx + 1].text.strip() if para_idx + 1 < len(doc.paragraphs) else ""
        if '由表' not in next_text and '测试' not in next_text and len(next_text) < 20:
            insert_paragraph_after(para, table_analysis[para_idx])
            print(f"  ✅ 表题段落{para_idx}后: 已补充结果分析")
        else:
            print(f"  ⚠️  段落{para_idx}后已有内容，跳过")

# ============================================================
# 10. 设置2.4.3允许西文在单词中间换行
# ============================================================
print("\n=== 10. 设置2.4.3西文换行 ===")
# 找到2.4.3天气查询节的段落
for i in range(177, 181):
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)
        # 设置允许西文在单词中间换行
        wordWrap = OxmlElement('w:wordWrap')
        wordWrap.set(qn('w:val'), '0')  # 0=允许在单词中间换行
        pPr.append(wordWrap)
        print(f"  ✅ 段落{i}: 已设置允许西文中间换行")

# ============================================================
# 保存
# ============================================================
output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
