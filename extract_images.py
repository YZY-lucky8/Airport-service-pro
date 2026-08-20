#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""提取海报文档中的图片，并分析报告章节结构"""
import os
from docx import Document

# 1. 提取图片
doc = Document(r'C:\Users\Lenovo\Desktop\附件3-路演海报要求(4).docx')
output_dir = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\poster_images'
os.makedirs(output_dir, exist_ok=True)

img_idx = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_idx += 1
        img_data = rel.target_part.blob
        ext = rel.target_ref.split('.')[-1]
        img_path = os.path.join(output_dir, f'poster_img_{img_idx}.{ext}')
        with open(img_path, 'wb') as f:
            f.write(img_data)
        print(f"提取图片{img_idx}: {img_path} ({len(img_data)} bytes)")

# 2. 分析报告章节结构
print("\n=== 最终版报告章节结构 ===")
doc2 = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_最终版.docx')
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    style = para.style.name
    # 只显示标题样式或编号章节
    if style.startswith('Heading') or (text and text[0].isdigit() and '.' in text[:5] and len(text) < 50):
        print(f"[{i}] ({style}) {text[:80]}")
