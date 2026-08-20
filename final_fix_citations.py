#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
综合修改脚本：
1. 修正问题1-5（布隆过滤器、第五章矛盾、内存等）
2. 插入17篇参考文献角标
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告.docx')

def replace_in_para(para, old_text, new_text):
    """在段落中替换文本，保留格式"""
    full_text = para.text
    if old_text not in full_text:
        return False
    # 简单方式：清空所有run，用第一个run的格式重写
    if para.runs:
        # 保存第一个run的格式
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        
        # 替换文本
        new_full = full_text.replace(old_text, new_text)
        
        # 清空所有run
        for run in para.runs:
            run.text = ''
        
        # 在第一个run写入新文本
        first_run.text = new_full
        return True
    return False

def add_citation_after_keyword(para, keyword, citation):
    """在关键词后插入参考文献角标"""
    full_text = para.text
    if keyword not in full_text:
        return False
    # 检查是否已有角标
    if citation in full_text:
        return False
    
    # 在关键词后插入角标
    new_text = full_text.replace(keyword, keyword + citation, 1)
    
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = new_text
        return True
    return False

# ============================================================
# 第一部分：修正问题1-5
# ============================================================
print("=== 修正问题1-5 ===")

# 问题1：段落132 - 布隆过滤器描述
para132 = doc.paragraphs[132]
old132 = "系统配置位数组大小为10000位，使用3个哈希函数（基于MurmurHash3变体实现），理论误判率约0.08%（低于0.1%的目标要求），内存占用约12KB/万条记录。"
new132 = "系统配置预期容量10000条、误判率0.001，基于公式动态计算得位数组大小143773位、10个哈希函数（采用双重哈希double hashing策略，由两个基础哈希函数派生k个位置），理论误判率约0.1%，内存占用约18KB/万条记录。"
if replace_in_para(para132, old132, new132):
    print("  ✅ 问题1: 段落132布隆过滤器描述已修正")
else:
    print("  ❌ 问题1: 段落132未找到匹配文本")
    print(f"  实际文本: {para132.text[:200]}")

# 问题2：段落526 - 第五章测试结果
para526 = doc.paragraphs[526]
if replace_in_para(para526, "功能测试 33 项中 26 项通过（78.8%）", "功能测试36项全部通过（100%）"):
    print("  ✅ 问题2: 段落526第五章测试结果已修正")
else:
    # 尝试不带空格的版本
    if replace_in_para(para526, "功能测试33项中26项通过（78.8%）", "功能测试36项全部通过（100%）"):
        print("  ✅ 问题2: 段落526第五章测试结果已修正(无空格版)")
    else:
        print("  ❌ 问题2: 段落526未找到匹配文本")
        print(f"  实际文本: {para526.text[:200]}")

# 问题3：段落114 - 内存70兆→77兆
para114 = doc.paragraphs[114]
if replace_in_para(para114, "约70兆字节", "约77兆字节"):
    print("  ✅ 问题3: 段落114内存已修正")
else:
    print("  ❌ 问题3: 段落114未找到匹配")

# 问题4：表22行3 - 布隆过滤器内存
table22 = doc.tables[22]
cell = table22.rows[3].cells[1]
if '12KB' in cell.text:
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run('≈18KB/万条')
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    print("  ✅ 问题4: 表22布隆过滤器内存已修正")
else:
    print(f"  ❌ 问题4: 表22未找到12KB，实际: {cell.text}")

# 问题5：段落206 - 阈值配置描述
para206 = doc.paragraphs[206]
fixed5 = False
if replace_in_para(para206, "布隆过滤器容量当前值10000位", "布隆过滤器预期容量当前值10000条"):
    fixed5 = True
if replace_in_para(para206, "建议值15000位", "建议值15000条"):
    fixed5 = True
if replace_in_para(para206, "布隆过滤器容量从5000扩容至10000", "布隆过滤器预期容量从5000条扩容至10000条"):
    fixed5 = True
if fixed5:
    print("  ✅ 问题5: 段落206阈值配置描述已修正")
else:
    print("  ❌ 问题5: 段落206未找到匹配")
    print(f"  实际文本: {para206.text[:200]}")

# ============================================================
# 第二部分：插入17篇参考文献角标
# ============================================================
print("\n=== 插入参考文献角标 ===")

# 定义角标插入规则：(段落索引, 关键词, 角标)
citations = [
    # [1] 关键信息基础设施安全保护条例
    (99, "关键信息基础设施安全保护条例", "[1]"),
    # [2] 等保2.0 - 在安全合规处添加
    (99, "要求，以极低成本", "[2]要求，以极低成本"),  # 特殊处理
    # [3] WCAG 2.1
    (83, "WCAG 2.1", "[3]"),
    # [4] JWT
    (146, "JWT令牌结构遵循标准格式", "[4]"),
    # [5] HMAC
    (135, "HMAC-SHA256签名算法", "[5]"),
    # [6] LangGraph
    (155, "LangGraph有向图状态机", "[6]"),
    # [7] 通义千问
    (166, "通义千问API", "[7]"),
    # [8] vLLM
    (166, "vLLM本地部署模式", "[8]"),
    # [9] OWASP Top 10 - 智能体安全风险
    (149, "多智能体防御流水线研究成果", "[9]"),
    # [10] CSRF
    (147, "CSRF Token机制", "[10]"),
    # [11] 恒定时间比较
    (147, "恒定时间比较防止时序攻击", "[11]"),
    # [12] 银河麒麟
    (113, "银河麒麟操作系统", "[12]"),
    # [13] 高德地图
    (185, "高德地图API", "[13]"),
    # [14] ECharts
    (113, "ECharts", "[14]"),
    # [15] Bloom Filter原始论文
    (77, "布隆过滤器（Bloom Filter）", "[15]"),
    # [16] DDoS攻击研究
    (65, "分布式拒绝服务攻击（DDoS）", "[16]"),
    # [17] DDoS防御研究
    (65, "导致终端服务不可用", "[17]"),
]

success_count = 0
for para_idx, keyword, citation in citations:
    if para_idx >= len(doc.paragraphs):
        print(f"  ❌ [{citation}] 段落{para_idx}不存在")
        continue
    para = doc.paragraphs[para_idx]
    
    # 特殊处理[2]
    if citation == "[2]要求，以极低成本":
        if replace_in_para(para, "要求，以极低成本", "要求[2]，以极低成本"):
            print(f"  ✅ [2] 已插入段落{para_idx}")
            success_count += 1
        else:
            print(f"  ❌ [2] 段落{para_idx}未找到匹配")
        continue
    
    if add_citation_after_keyword(para, keyword, citation):
        print(f"  ✅ {citation} 已插入段落{para_idx} ('{keyword}')")
        success_count += 1
    else:
        print(f"  ❌ {citation} 段落{para_idx}未找到关键词'{keyword}'或已存在")
        print(f"     实际文本: {para.text[:100]}")

print(f"\n角标插入完成: {success_count}/17")

# ============================================================
# 保存
# ============================================================
output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_最终版.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
