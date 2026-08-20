#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""补充修改：验证表0 + 为4.1.2添加对比优势段落"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import copy

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版.docx')

print("=== 验证表0（1.2.2节）===")
table0 = doc.tables[0]
for ri, row in enumerate(table0.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{ri}: {cells}")

# 为4.1.2添加对比优势段落
print("\n=== 为4.1.2添加对比优势 ===")
compare_text = ('对比优势：相比传统硬件防火墙（性能损耗>30%、单台数万元）与云端WAF'
                '（网络延迟50-200ms、依赖外网连接），本方案以7.1%的极低性能损耗实现六层纵深防御，'
                '稳定运行内存仅约77MB，且支持完全离线部署，满足机场内网隔离与自主可控合规要求。')

for i, para in enumerate(doc.paragraphs):
    if '4.1.2 资源受限终端的轻量化纵深防御理念' in para.text:
        # 找到4.1.2的具体体现段落（应该是i+2）
        target_idx = i + 2
        if target_idx < len(doc.paragraphs):
            target_para = doc.paragraphs[target_idx]
            # 在具体体现后插入对比优势
            new_p = copy.deepcopy(target_para._element)
            for child in list(new_p):
                if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                    new_p.remove(child)
            target_para._element.addnext(new_p)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, target_para._parent)
            run = new_para.add_run(compare_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            print(f"  在段落{target_idx}后插入了4.1.2对比优势")
        break

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版.docx')
print("\n✅ 保存完成")
