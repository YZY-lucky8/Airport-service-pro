#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""补充修复：表3-8分析、标点修正、页码检查"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

def insert_paragraph_after(para, text):
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    new_para = Paragraph(new_p, doc)
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return new_para

from docx.shared import Pt

# 1. 修正图2-9引用的标点错误
print("=== 修正图2-9引用标点 ===")
for i, para in enumerate(doc.paragraphs):
    if '三道防线。，整体运行机制' in para.text:
        set_para_text(para, para.text.replace('三道防线。，整体运行机制', '三道防线，整体运行机制'))
        print(f"  ✅ 段落{i}: 标点已修正")
        break

# 2. 补充表3-8后的结果分析
print("\n=== 补充表3-8结果分析 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('表 3-8'):
        # 检查下一段
        next_text = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ""
        if '3.7' in next_text or '由表' not in next_text:
            insert_paragraph_after(para, "由表3-8可知，三道智能体安全防线在输入过滤、运行时监控、输出过滤各阶段的测试场景中均正确拦截恶意请求，无漏防误报，智能体安全防护能力达标。")
            print(f"  ✅ 表3-8后: 已补充结果分析")
        break

# 3. 补充表3-10后的结果分析
print("\n=== 补充表3-10结果分析 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('表 3-10'):
        next_text = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ""
        if '3.7.3' in next_text or '由表' not in next_text:
            insert_paragraph_after(para, "由表3-10可知，系统在启动、稳定运行、压测峰值各阶段的内存占用均在设计目标内，CPU使用率合理，资源占用表现优异。")
            print(f"  ✅ 表3-10后: 已补充结果分析")
        break

# 4. 检查并添加页码
print("\n=== 检查页码设置 ===")
for si, section in enumerate(doc.sections):
    footer = section.footer
    has_page_num = False
    for para in footer.paragraphs:
        # 检查是否有页码域
        for run in para.runs:
            if 'PAGE' in run._element.xml or 'page' in run._element.xml.lower():
                has_page_num = True
                break
    print(f"  节{si}: 页脚有页码域: {has_page_num}")
    
    # 如果没有页码，添加页码
    if not has_page_num and footer.paragraphs:
        para = footer.paragraphs[0]
        para.alignment = 1  # 居中
        # 添加页码域
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        print(f"  ✅ 节{si}: 已添加页码")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')
print("\n✅ 已保存")
