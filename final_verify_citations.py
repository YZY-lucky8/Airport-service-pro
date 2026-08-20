#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""给图3-3补充引用，并正确验证"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

# 给图3-3 unifyResponse.js添加引用 - 在攻击反馈隐藏测试描述中
print("=== 给图3-3补充引用 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '攻击反馈隐藏' in text and '如图3-3' not in text and not re.match(r'^图\s', text):
        new_text = text + "（如图3-3所示）"
        set_para_text(para, new_text)
        print(f"  ✅ 段落{i}: 添加如图3-3所示")
        print(f"     原文: {text[:60]}")
        break

# 正确验证
print("\n=== 最终验证 ===")
# 收集所有图题编号
fig_captions = set()
table_captions = set()
for para in doc.paragraphs:
    text = para.text.strip()
    m = re.match(r'^图\s+(\d+-\d+)', text)
    if m: fig_captions.add(m.group(1))
    m = re.match(r'^表\s+(\d+-\d+)', text)
    if m: table_captions.add(m.group(1))

# 收集所有正文引用编号
fig_cited = set()
table_cited = set()
for para in doc.paragraphs:
    text = para.text
    for m in re.findall(r'如图\s*(\d+-\d+)', text):
        fig_cited.add(m)
    for m in re.findall(r'如表\s*(\d+-\d+)', text):
        table_cited.add(m)

print(f"图题: {len(fig_captions)}个, 已引用: {len(fig_cited)}个")
missing_fig = fig_captions - fig_cited
if missing_fig:
    print(f"  缺失: {missing_fig}")
else:
    print("  ✅ 全部有引用")

print(f"表题: {len(table_captions)}个, 已引用: {len(table_cited)}个")
missing_table = table_captions - table_cited
if missing_table:
    print(f"  缺失: {missing_table}")
else:
    print("  ✅ 全部有引用")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')
print("\n✅ 已保存")
