#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""最终验证"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx')

print("=== 验证结果 ===")

# 1. 括号模式引用
bracket = 0
for para in doc.paragraphs:
    if re.search(r'（如[图表]\d', para.text):
        bracket += 1
print(f"1. 剩余括号模式引用: {bracket} 处 {'✅' if bracket==0 else '❌'}")

# 2. 图3-1引用
fig3_cited = False
for para in doc.paragraphs:
    if '如图3-1' in para.text and not re.match(r'^图\s+3', para.text.strip()):
        fig3_cited = True
        break
print(f"2. 图3-1有正文引用: {'✅' if fig3_cited else '❌'}")

# 3. 表3-7说明
fig37 = False
for para in doc.paragraphs:
    if '第一级精确匹配命中' in para.text:
        fig37 = True
        break
print(f"3. 表3-7有说明: {'✅' if fig37 else '❌'}")

# 4. 3.2.1有介绍
sec321 = False
for para in doc.paragraphs:
    if '本节介绍系统测试环境的具体配置' in para.text:
        sec321 = True
        break
print(f"4. 3.2.1有介绍: {'✅' if sec321 else '❌'}")

# 5. 3.4.2有介绍
sec342 = False
for para in doc.paragraphs:
    if '本节介绍情感分析模块的测试方法' in para.text:
        sec342 = True
        break
print(f"5. 3.4.2有介绍: {'✅' if sec342 else '❌'}")

# 6. 表4-1有介绍
tab41 = False
for para in doc.paragraphs:
    if '本系统所提方案与传统商业WAF' in para.text:
        tab41 = True
        break
print(f"6. 表4-1有介绍: {'✅' if tab41 else '❌'}")

# 7. 小圆圈占位符
circle = 0
for para in doc.paragraphs:
    if '⃝' in para.text:
        circle += 1
print(f"7. 小圆圈占位符: {circle} 处 {'✅' if circle==0 else '❌'}")

# 8. 表3-8有分析
tab38 = False
for para in doc.paragraphs:
    if '由表3-8可知' in para.text:
        tab38 = True
        break
print(f"8. 表3-8有结果分析: {'✅' if tab38 else '❌'}")

print("\n✅ 全部验证完成")
