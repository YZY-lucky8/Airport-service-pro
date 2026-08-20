#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""修正引用文字重复问题，并检查遗漏的引用"""
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')

# 第一步：修正所有"如图图"和"如表表"的重复
print("=== 修正引用文字重复 ===")
fix_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '如图图' in text or '如表表' in text:
        new_text = text.replace('如图图', '如图').replace('如表表', '如表')
        # 设置文本
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs:
                run.text = ''
            first_run.text = new_text
        fix_count += 1
        print(f"  段落{i}: 修正引用文字")

print(f"共修正 {fix_count} 处")

# 第二步：检查哪些图表还没有正文引用
print("\n=== 检查图表引用覆盖情况 ===")

# 收集所有图题和表题
all_figs = []
all_tables = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^图\s+\d+-\d+', text):
        fig_id = re.match(r'^图\s+(\d+-\d+)', text).group(1)
        all_figs.append((i, fig_id, text[:50]))
    elif re.match(r'^表\s+\d+-\d+', text):
        tbl_id = re.match(r'^表\s+(\d+-\d+)', text).group(1)
        all_tables.append((i, tbl_id, text[:50]))

# 收集正文中所有引用
all_citations = set()
for para in doc.paragraphs:
    text = para.text
    # 匹配"如图X-Y所示"或"如表X-Y所示"
    fig_matches = re.findall(r'如图\s*(\d+-\d+)', text)
    tbl_matches = re.findall(r'如表\s*(\d+-\d+)', text)
    for m in fig_matches:
        all_citations.add(('fig', m))
    for m in tbl_matches:
        all_citations.add(('tbl', m))

print(f"图题总数: {len(all_figs)}")
print(f"表题总数: {len(all_tables)}")
print(f"正文引用数: {len(all_citations)}")

# 检查缺失的引用
missing_figs = []
for idx, fig_id, text in all_figs:
    if ('fig', fig_id) not in all_citations:
        missing_figs.append((idx, fig_id, text))

missing_tables = []
for idx, tbl_id, text in all_tables:
    if ('tbl', tbl_id) not in all_citations:
        missing_tables.append((idx, tbl_id, text))

if missing_figs:
    print(f"\n缺失引用的图 ({len(missing_figs)}个):")
    for idx, fid, text in missing_figs:
        print(f"  图{fid} (段落{idx}): {text}")

if missing_tables:
    print(f"\n缺失引用的表 ({len(missing_tables)}个):")
    for idx, tid, text in missing_tables:
        print(f"  表{tid} (段落{idx}): {text}")

if not missing_figs and not missing_tables:
    print("\n✅ 所有图表都有正文引用！")

# 保存
doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')
print("\n✅ 已保存修正版")
