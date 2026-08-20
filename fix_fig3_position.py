#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""修复图3-1引用位置：删除段落112的错误引用，在3.6.1节正确添加"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

# 1. 删除段落112中错误添加的引用
print("=== 修复图3-1引用位置 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '边缘层部署于Node.js' in text and '图3-1' in text:
        # 删除错误添加的引用部分
        new_text = re.sub(r'\s*相关测试结果分别如图3-1、图3-2、图3-3所示。', '', text)
        set_para_text(para, new_text)
        print(f"  ✅ 段落{i}: 已删除错误引用")
        break

# 2. 在3.6.1节"白名单与关键服务免扰"段落正确添加
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text.strip().startswith('白名单与关键服务免扰') and '图3' not in text:
        set_para_text(para, text + " 相关测试结果分别如图3-1、图3-2、图3-3所示。")
        print(f"  ✅ 段落{i}: 已正确添加图3-1/3-2/3-3引用")
        break

# 验证
print("\n=== 验证 ===")
for i, para in enumerate(doc.paragraphs):
    if '图3-1' in para.text and not re.match(r'^图\s+3', para.text.strip()):
        print(f"  引用位置: 段落{i} - {para.text[:60]}...")

output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx'
doc.save(output)
print(f"\n✅ 已保存")
