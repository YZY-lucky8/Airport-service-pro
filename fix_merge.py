#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""综合修改：图表引用优化、表格前后文字补充、格式修复"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\作品报告-(3) - 用于合并.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

def insert_paragraph_after(para, text):
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    new_para = Paragraph(new_p, doc)
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return new_para

def insert_paragraph_before(para, text):
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, doc)
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return new_para

def set_table_header_repeat(table):
    tr = table.rows[0]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

# ============================================================
# 1. 优化图表引用方式（20处）
# ============================================================
print("=== 1. 优化图表引用方式 ===")
citation_fixes = [
    (44, "系统各项核心指标均达到或超过设计目标：（如表1-1所示）", "系统各项核心指标均达到或超过设计目标，如表1-1所示："),
    (81, "本项目目标客户聚焦于公共服务场景运营方与渠道合作方：（如表1-2所示）", "本项目目标客户聚焦于公共服务场景运营方与渠道合作方，如表1-2所示："),
    (124, "（一）滑动窗口频率检测。（如图2-4所示）该机制采用", "（一）滑动窗口频率检测。如图2-4所示，该机制采用"),
    (127, "布隆过滤器IP黑名单。（如图2-5所示）被滑动窗口检测拦截的IP地址自动加入黑名单", "布隆过滤器IP黑名单。图2-5展示被滑动窗口检测拦截的IP地址自动加入黑名单的机制，后续请求通过布隆过滤器进行微秒级查询拦截"),
    (131, "（三）HMAC一次性令牌。（如图2-6所示）该机制基于", "（三）HMAC一次性令牌。如图2-6所示，该机制基于"),
    (136, "（六）关键服务免扰。（如图2-7所示） 紧急帮助接口", "（六）关键服务免扰。如图2-7所示，紧急帮助接口"),
    (139, "请求成功率超过99.2%。（如图2-8所示）", "请求成功率超过99.2%，如图2-8所示。"),
    (209, "各表存储用途、核心字段、预置数据量整理如下表：（如表2-1所示）", "各表存储用途、核心字段、预置数据量整理如下，如表2-1所示："),
    (227, "实测，系统核心性能指标如下：（如表2-3所示）", "实测，系统核心性能指标如下，如表2-3所示："),
    (262, "知识库实际包含三十三条条目，覆盖七分类：（如表3-2所示）", "知识库实际包含三十三条条目，覆盖七分类，如表3-2所示："),
    (281, "航班查询测试小结：（如表3-3所示）", "航班查询测试小结如表3-3所示："),
    (308, "知识库问答测试小结：（如表3-4所示）", "知识库问答测试小结如表3-4所示："),
    (363, "系统定义了九类预定义意图，经实际测试验证：（如表3-5所示）", "系统定义了九类预定义意图，经实际测试验证，结果如表3-5所示："),
    (370, "双字符分词处理中文文本。（如表3-7所示）", "双字符分词处理中文文本，测试结果如表3-7所示："),
    (419, "通过接口响应 latency 字段实测：（如表3-9所示）", "通过接口响应 latency 字段实测，结果如表3-9所示："),
    (474, '形成"先放行特殊→再检查白名单→后逐层过滤"的优化路径：（如图4-1）', '形成"先放行特殊→再检查白名单→后逐层过滤"的优化路径，如图4-1所示。'),
    (478, "对比优势：（如表4-1所示）", "本系统所提方案与传统商业WAF、单机制轻量方案的对比如表4-1所示："),
    (485, "对比优势：（如表4-2所示）", "本系统三道防线方案与传统输入过滤方案、VeriGuard形式化验证方案的对比如表4-2所示："),
    (492, "对比优势：（如表4-3所示）", "本系统三级检索方案与向量数据库RAG、简单关键词匹配方案的对比如表4-3所示："),
    (500, "对比优势：（如表4-4所示）", "本系统自研框架与LangChain/LangGraph、CrewAI、AutoGen等主流框架的对比如表4-4所示："),
]

for para_idx, old, new in citation_fixes:
    if para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        if old in para.text:
            set_para_text(para, para.text.replace(old, new))
            print(f"  ✅ 段落{para_idx}")
        else:
            print(f"  ⚠️  段落{para_idx}未匹配")

# ============================================================
# 2. 跨页表格设置重复标题行
# ============================================================
print("\n=== 2. 设置跨页表格重复标题行 ===")
for ti in [2, 4, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 19, 20, 21]:
    if ti < len(doc.tables):
        set_table_header_repeat(doc.tables[ti])
print(f"  ✅ 15个表格已设置重复标题行")

# ============================================================
# 3. 补充只有表格无内容的章节
# ============================================================
print("\n=== 3. 补充表格前导文字 ===")
# 3.2.1 测试环境配置
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '3.2.1 测试环境配置':
        insert_paragraph_after(para, "本节介绍系统测试环境的具体配置，包括服务器硬件、操作系统、运行时环境以及数据库等，确保测试结果可复现，具体配置如表3-1所示。")
        print("  ✅ 3.2.1测试环境介绍")
        break

# 3.4.2 情感分析模块测试
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '3.4.2 情感分析模块测试':
        insert_paragraph_after(para, "本节介绍情感分析模块的测试方法，通过输入带有一定感情色彩的日常对话文本，系统进行情感预判，然后与测试样本的标注标签对比，验证情感识别的准确性，测试结果如表3-6所示。")
        print("  ✅ 3.4.2情感分析介绍")
        break

# ============================================================
# 4. 表3-7补充说明
# ============================================================
print("\n=== 4. 表3-7补充说明 ===")
for i, para in enumerate(doc.paragraphs):
    if '检索引擎工作稳定' in para.text and '第一级精确匹配' in para.text:
        set_para_text(para, para.text + " 表中所有测试用例均在第一级精确匹配命中，原因在于测试用例均使用知识库条目标题中的标准关键词进行查询，符合三级检索引擎的设计预期——标准查询优先在第一级精确匹配，仅当精确匹配无结果时才降级至第二级模糊匹配和第三级全文检索。")
        print("  ✅ 表3-7后补充说明")
        break

# ============================================================
# 5. 图3-1到3-3补充引用
# ============================================================
print("\n=== 5. 图3-1到3-3补充引用 ===")
for i, para in enumerate(doc.paragraphs):
    if '白名单与关键服务免扰' in para.text and '如图3' not in para.text:
        set_para_text(para, para.text + " 相关测试结果分别如图3-1、图3-2、图3-3所示。")
        print(f"  ✅ 段落{i}补充图3-1/3-2/3-3引用")
        break

# ============================================================
# 6. 测试表格后补充结果分析
# ============================================================
print("\n=== 6. 补充测试表格结果分析 ===")
table_analysis = {
    '表 3-8': "由表3-8可知，三道智能体安全防线在输入过滤、运行时监控、输出过滤各阶段的测试场景中均正确拦截恶意请求，无漏防误报，智能体安全防护能力达标。",
    '表 3-9': "由表3-9可知，智能体规则引擎在各类业务意图下处理延迟均低于5ms，平均延迟约1ms，远优于≤10ms的设计目标。",
    '表 3-10': "由表3-10可知，系统在启动、稳定运行、压测峰值各阶段的内存占用均在设计目标内，CPU使用率合理，资源占用表现优异。",
}

# 从后往前插入避免索引变化
for caption, analysis in sorted(table_analysis.items(), reverse=True):
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(caption):
            next_text = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ""
            if '由表' not in next_text and len(next_text) < 20:
                insert_paragraph_after(para, analysis)
                print(f"  ✅ {caption}后补充分析")
            break

# ============================================================
# 7. 表2-2第三列表头
# ============================================================
print("\n=== 7. 表2-2第三列表头 ===")
table22 = doc.tables[3]
cell = table22.rows[0].cells[2]
if not cell.text.strip():
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run('核心字段与说明')
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    print("  ✅ 表2-2第三列表头已添加")

# ============================================================
# 8. 设置不压缩图片
# ============================================================
print("\n=== 8. 设置不压缩图片 ===")
settings = doc.settings.element
if settings.find(qn('w:doNotCompressPictures')) is None:
    doNotCompress = OxmlElement('w:doNotCompressPictures')
    settings.append(doNotCompress)
    print("  ✅ 已设置不压缩图片")

# ============================================================
# 保存
# ============================================================
output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
