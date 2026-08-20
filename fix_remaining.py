#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""检查并修复剩余问题"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

print("=== 检查1.2.2节备注 ===")
in_122 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1.2.2' in text:
        in_122 = True
    if in_122 and ('注：' in text or '备注：' in text):
        print(f"  [{i}] 仍存在: {text[:80]}")
        # 删除
        para._element.getparent().remove(para._element)
        print(f"  ✅ 已删除")
    if in_122 and text.startswith('1.3'):
        break

print("\n=== 检查剩余括号模式引用 ===")
bracket_count = 0
for i, para in enumerate(doc.paragraphs):
    if re.search(r'（如[图表]\d+-\d+所示）', para.text):
        bracket_count += 1
        print(f"  [{i}] {para.text[:80]}")
print(f"剩余括号模式引用: {bracket_count} 处")

print("\n=== 检查小圆圈占位符 ===")
circle_count = 0
for i, para in enumerate(doc.paragraphs):
    if '⃝' in para.text or '○' in para.text:
        circle_count += 1
        print(f"  [{i}] {para.text[:60]}")
print(f"占位符数量: {circle_count}")

# 设置文档不压缩图片
print("\n=== 设置文档不压缩图片 ===")
# 在文档设置中添加不压缩图片选项
settings = doc.settings.element
# 检查是否已有doNotEmbedSystemFonts等
# 添加不压缩图片
existing = settings.find(qn('w:doNotCompressPictures'))
if existing is None:
    doNotCompress = OxmlElement('w:doNotCompressPictures')
    settings.append(doNotCompress)
    print("  ✅ 已设置不压缩图片")
else:
    print("  已存在不压缩图片设置")

# 设置默认高保真分辨率
# 这个需要在Word选项中设置，docx中可以尝试设置
print("\n=== 验证表格结果分析 ===")
# 检查几个关键表格后是否有分析
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('由表3-') and len(text) > 30:
        print(f"  [{i}] {text[:60]}...")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')
print("\n✅ 已保存")
