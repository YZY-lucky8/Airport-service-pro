#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版_v3.docx')
table4 = doc.tables[4]

def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 行5: Agent规则引擎延迟
set_cell(table4.rows[5].cells[2], '<5ms（实测平均1ms）')
print("行5已更新: Agent规则引擎延迟")

# 行9: API接口测试覆盖率
set_cell(table4.rows[9].cells[2], '核心接口100%（总接口78.4%）')
print("行9已更新: API接口测试覆盖率")

# 行10: 语音识别准确率
set_cell(table4.rows[10].cells[2], '86.3%（Web Speech API标称）')
print("行10已更新: 语音识别准确率")

doc.save(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版_v3.docx')
print("\n✅ 保存完成")
