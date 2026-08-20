#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""精确定位所有修改点和参考文献插入位置"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告.docx')

# 需要查看的段落索引
target_paras = [77, 114, 132, 145, 146, 147, 149, 155, 165, 166, 185, 206, 225, 232, 448, 453, 476, 490, 496, 526]

print("=== 关键段落完整文本 ===")
for i in target_paras:
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if text:
            print(f"\n[{i}] {text}")

# 查看表22
print("\n\n=== 表22完整内容 ===")
if len(doc.tables) > 22:
    for ri, row in enumerate(doc.tables[22].rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{ri}: {cells}")

# 查找适合插入参考文献的关键段落
print("\n\n=== 参考文献插入候选位置 ===")
candidates = {
    '关键信息基础设施': [1],
    '等级保护': [2],
    '适老化': [3],
    '无障碍': [3],
    'JWT': [4],
    'HMAC': [5],
    'LangGraph': [6],
    '通义': [7],
    'vLLM': [8],
    'OWASP': [9, 10],
    'timingSafeEqual': [11],
    '恒定时间': [11],
    '银河麒麟': [12],
    '高德': [13],
    'ECharts': [14],
    '布隆过滤器': [15],
    'Bloom Filter': [15],
    'DDoS': [16, 17],
    '拒绝服务': [16, 17],
}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text or len(text) < 20:
        continue
    for kw, refs in candidates.items():
        if kw in text:
            # 检查是否已有角标
            import re
            if not re.search(r'\[\d+\]', text):
                print(f"[{i}] 关键词'{kw}'→{refs}: {text[:80]}...")
                break
