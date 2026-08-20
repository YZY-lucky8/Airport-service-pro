#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
按章节重新编号所有图、表，补全表题，添加正文引用
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
import copy

doc = Document(r'C:\Users\Lenovo\Desktop\作品报告-08.15-2.docx')

# 章节定义
chapters = []
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading 1') and para.text.strip():
        chapters.append((para._element, para.text.strip()))

def get_chapter(element):
    """获取元素所属章节"""
    chap = "前言"
    for chap_elem, chap_name in chapters:
        # 检查chap_elem是否在element之前
        if element is chap_elem:
            return chap_name
    # 遍历body子元素
    body = doc.element.body
    found_chap = "前言"
    for child in body.iterchildren():
        if child is element:
            return found_chap
        for chap_elem, chap_name in chapters:
            if child is chap_elem:
                found_chap = chap_name
                break
    return found_chap

def get_chapter_num(chap_name):
    """从章节名获取章节号"""
    if '第一章' in chap_name or '作品概述' in chap_name:
        return 1
    elif '第二章' in chap_name or '设计与实现' in chap_name:
        return 2
    elif '第三章' in chap_name or '测试与分析' in chap_name:
        return 3
    elif '创新' in chap_name:
        return 4
    elif '第五章' in chap_name or '总结' in chap_name:
        return 5
    return 0

def insert_paragraph_before(element, text, style_name=None):
    """在指定元素前插入段落"""
    new_p = OxmlElement('w:p')
    element.addprevious(new_p)
    para = Paragraph(new_p, doc)
    if text:
        run = para.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if style_name:
        para.style = doc.styles[style_name]
    return para

def set_para_text(para, text):
    """设置段落文本，保留第一个run的格式"""
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

# ============================================================
# 第一步：收集所有图题和表格的位置信息
# ============================================================
print("=== 收集图表位置 ===")
body = doc.element.body
children = list(body.iterchildren())

figures = []  # (element, old_text, chapter_num)
tables = []   # (element, has_caption, caption_element, chapter_num)

current_chap = "前言"
for idx, child in enumerate(children):
    if child.tag == qn('w:p'):
        para = Paragraph(child, doc)
        text = para.text.strip()
        # 更新章节
        for chap_elem, chap_name in chapters:
            if child is chap_elem:
                current_chap = chap_name
                break
        # 检查是否是图题
        if re.match(r'^图\s*\d+', text):
            chap_num = get_chapter_num(current_chap)
            figures.append((child, text, chap_num, idx))
    elif child.tag == qn('w:tbl'):
        chap_num = get_chapter_num(current_chap)
        # 检查前面是否有表题段落
        has_caption = False
        caption_elem = None
        for j in range(idx-1, max(0, idx-3), -1):
            if children[j].tag == qn('w:p'):
                p = Paragraph(children[j], doc)
                if re.match(r'^表\s*\d+', p.text.strip()):
                    has_caption = True
                    caption_elem = children[j]
                    break
                elif p.text.strip():
                    break
        tables.append((child, has_caption, caption_elem, chap_num, idx))

print(f"找到 {len(figures)} 个图题, {len(tables)} 个表格")

# ============================================================
# 第二步：按章节重新编号图题
# ============================================================
print("\n=== 重新编号图题 ===")
fig_counters = {}  # chap_num -> count
fig_new_names = {}  # element -> new_name

for elem, old_text, chap_num, idx in figures:
    if chap_num not in fig_counters:
        fig_counters[chap_num] = 0
    fig_counters[chap_num] += 1
    new_num = fig_counters[chap_num]
    
    # 提取原图题描述部分
    # 格式："图 X 描述" 或 "图X 描述"
    desc = re.sub(r'^图\s*\d+\s*', '', old_text)
    new_text = f"图 {chap_num}-{new_num} {desc}"
    
    para = Paragraph(elem, doc)
    set_para_text(para, new_text)
    fig_new_names[elem] = new_text
    print(f"  章{chap_num}: '{old_text[:40]}' -> '{new_text[:40]}'")

# ============================================================
# 第三步：给表格添加/修改表题
# ============================================================
print("\n=== 处理表题 ===")
table_counters = {}  # chap_num -> count
table_new_names = {}  # element -> new_name

# 预定义每个表格的表题（按文档流顺序）
# 我们需要根据表格内容来确定表题
table_captions = {
    # 第一章
    0: "系统核心性能指标概览表",
    1: "目标用户群体与方案价值分析表",
    # 第二章
    2: "核心数据表清单",
    3: "智能体扩展数据表清单",
    4: "系统核心性能指标表",
    # 第三章
    5: "测试环境配置表",
    6: "知识库分类与优先级统计表",
    7: "航班查询功能测试结果表",
    8: "知识库问答功能测试结果表",
    9: "意图分类引擎测试结果表",
    10: "情感分析模块测试结果表",
    11: "三级混合检索引擎测试结果表",
    12: "智能体安全防线测试结果表",
    13: "智能体处理延迟测试结果表",
    14: "系统资源占用测试表",
    15: "WCAG 2.1第二级合规性验证表",
    16: "系统交互总体数据表",
    17: "用户意图分布统计表",
    18: "知识库分类统计表",
    19: "功能测试汇总表",
    20: "安全防护测试汇总表",
    21: "性能指标汇总表",
    # 第四章
    22: "DDoS防护方案对比表",
    23: "智能体安全防线方案对比表",
    24: "RAG检索方案对比表",
    25: "主流智能体框架对比表",
}

for ti, (elem, has_caption, caption_elem, chap_num, idx) in enumerate(tables):
    if chap_num not in table_counters:
        table_counters[chap_num] = 0
    table_counters[chap_num] += 1
    new_num = table_counters[chap_num]
    
    caption_text = table_captions.get(ti, f"表格{ti+1}")
    new_caption = f"表 {chap_num}-{new_num} {caption_text}"
    
    if has_caption and caption_elem is not None:
        # 修改已有表题
        para = Paragraph(caption_elem, doc)
        set_para_text(para, new_caption)
        print(f"  表{ti+1}(章{chap_num}): 修改表题 -> '{new_caption}'")
    else:
        # 在表格前插入新表题
        new_para = insert_paragraph_before(elem, new_caption)
        # 设置表题格式：居中、宋体、小五号
        new_para.alignment = 1  # 居中
        for run in new_para.runs:
            run.font.size = Pt(9)
        print(f"  表{ti+1}(章{chap_num}): 插入表题 -> '{new_caption}'")
    
    table_new_names[elem] = new_caption

# ============================================================
# 第四步：在正文中添加图表引用
# ============================================================
print("\n=== 添加正文引用 ===")

# 重新遍历文档流，在每个图题/表题前面的描述段落末尾添加引用
children = list(body.iterchildren())
current_chap = "前言"
citation_count = 0

for idx, child in enumerate(children):
    if child.tag == qn('w:p'):
        para = Paragraph(child, doc)
        text = para.text.strip()
        for chap_elem, chap_name in chapters:
            if child is chap_elem:
                current_chap = chap_name
                break
        
        # 检查是否是图题
        if re.match(r'^图\s+\d+-\d+', text):
            # 找前面最近的非空、非标题、非图题的描述段落
            for j in range(idx-1, max(0, idx-8), -1):
                if children[j].tag == qn('w:p'):
                    p = Paragraph(children[j], doc)
                    pt = p.text.strip()
                    if pt and not re.match(r'^图\s', pt) and not re.match(r'^表\s', pt) and not p.style.name.startswith('Heading'):
                        # 在段落末尾添加引用
                        if '如图' not in pt and '图' + text.split()[1] not in pt:
                            fig_ref = text.split()[0] + text.split()[1]  # 图X-Y
                            # 检查是否已有引用
                            if f'如图{fig_ref}' not in pt:
                                new_text = pt + f"（如图{fig_ref}所示）"
                                set_para_text(p, new_text)
                                citation_count += 1
                                print(f"  图引用: 段落{j}添加'如图{fig_ref}所示'")
                        break
        
        # 检查是否是表题
        if re.match(r'^表\s+\d+-\d+', text):
            # 找前面最近的非空、非标题、非表题的描述段落
            for j in range(idx-1, max(0, idx-8), -1):
                if children[j].tag == qn('w:p'):
                    p = Paragraph(children[j], doc)
                    pt = p.text.strip()
                    if pt and not re.match(r'^表\s', pt) and not re.match(r'^图\s', pt) and not p.style.name.startswith('Heading'):
                        if '如表' not in pt:
                            tbl_ref = text.split()[0] + text.split()[1]  # 表X-Y
                            if f'如表{tbl_ref}' not in pt:
                                new_text = pt + f"（如表{tbl_ref}所示）"
                                set_para_text(p, new_text)
                                citation_count += 1
                                print(f"  表引用: 段落{j}添加'如表{tbl_ref}所示'")
                        break

print(f"\n共添加 {citation_count} 处正文引用")

# ============================================================
# 保存
# ============================================================
output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
