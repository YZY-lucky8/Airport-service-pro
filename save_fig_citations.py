#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""重新加载并保存（换文件名）"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

def set_para_text(para, text):
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = text
    else:
        para.text = text

def insert_paragraph_before(para, text):
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, doc)
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return new_para

# 检查是否已经添加了引用（避免重复）
has_2_1 = False
has_2_3 = False
has_2_10 = False
for para in doc.paragraphs:
    if '系统整体架构如图2-1所示' in para.text:
        has_2_1 = True
    if '态势总览页面如图2-3所示' in para.text:
        has_2_3 = True
    if '主界面地图导航页如图2-10所示' in para.text:
        has_2_10 = True

print(f"已存在引用: 图2-1={has_2_1}, 图2-3={has_2_3}, 图2-10={has_2_10}")

if not has_2_1:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('图 2-1'):
            insert_paragraph_before(para, "本章详细阐述系统的设计方案与实现细节。系统整体架构如图2-1所示，采用分层设计思想，将安全防护与智能服务深度融合，确保各模块职责清晰、耦合度低。")
            print("✅ 图2-1引用已添加")
            break

if not has_2_3:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('图 2-3'):
            insert_paragraph_before(para, "安全防护模块是系统的核心组成部分，机场安全运营中心态势总览页面如图2-3所示，实时展示系统运行状态、攻击告警、防护效果等关键信息。")
            print("✅ 图2-3引用已添加")
            break

if not has_2_10:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('图 2-10'):
            insert_paragraph_before(para, "核心功能模块为旅客提供智能化出行服务，主界面地图导航页如图2-10所示，集成了航班查询、地点导航、天气查询、知识库问答等功能入口。")
            print("✅ 图2-10引用已添加")
            break

# 图2-2和图2-11在正文中添加引用
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '三级轻量化架构' in text and '如图2-2' not in text:
        set_para_text(para, text.replace('三级轻量化架构，层次清晰、分工明确。', '三级轻量化架构，层次清晰、分工明确，系统总体架构与技术路线如图2-2所示。'))
        print("✅ 图2-2引用已添加")
    if '可视化机舱布局选座界面' in text and '如图2-11' not in text:
        set_para_text(para, text.replace('加载可视化机舱布局选座界面', '加载可视化机舱布局选座界面，旅客端自助选座界面如图2-11所示'))
        print("✅ 图2-11引用已添加")

# 验证
fig_cited = set()
for para in doc.paragraphs:
    text = para.text
    if not re.match(r'^图\s+\d+-\d+', text.strip()):
        for m in re.findall(r'图\s*(\d+-\d+)', text):
            fig_cited.add(m)

all_figs = set()
for para in doc.paragraphs:
    m = re.match(r'^图\s+(\d+-\d+)', para.text.strip())
    if m:
        all_figs.add(m.group(1))

missing = all_figs - fig_cited
print(f"\n图总数: {len(all_figs)}, 已引用: {len(fig_cited)}, 缺失: {missing}")

output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版_v2.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
