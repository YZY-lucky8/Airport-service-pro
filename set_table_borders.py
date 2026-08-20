#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将所有表格边框设置为黑色实线"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx')

def set_table_borders(table, color="000000", size="4"):
    """
    设置表格边框为黑色实线
    color: 十六进制颜色，000000为黑色
    size: 边框粗细，单位1/8磅，4=0.5磅，8=1磅
    """
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 移除已有的边框设置
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    
    # 创建新的边框设置
    tblBorders = OxmlElement('w:tblBorders')
    
    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 实线
        border.set(qn('w:sz'), size)       # 粗细
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)   # 黑色
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# 处理所有表格
print(f"=== 设置 {len(doc.tables)} 个表格边框为黑色实线 ===")
for i, table in enumerate(doc.tables):
    set_table_borders(table, color="000000", size="6")  # 6 = 0.75磅
    print(f"  表{i+1}: 已设置黑色实线边框")

# 验证表题
print("\n=== 验证表题 ===")
import re
table_captions = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if re.match(r'^表\s+\d+-\d+', text):
        table_captions += 1

print(f"表题数量: {table_captions}")
print(f"表格数量: {len(doc.tables)}")
if table_captions >= len(doc.tables):
    print("✅ 所有表格都有表题")
else:
    print(f"⚠️  缺少 {len(doc.tables) - table_captions} 个表题")

# 保存
output = r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_图表编号修正版.docx'
doc.save(output)
print(f"\n✅ 已保存至: {output}")
