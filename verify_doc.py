#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""验证修改后的文档"""
from docx import Document

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版.docx')

print("=== 验证性能指标表4（2.7.4节）===")
table4 = doc.tables[4]
for ri, row in enumerate(table4.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{ri}: {cells}")

print("\n=== 验证部署备注 ===")
for i, para in enumerate(doc.paragraphs):
    if '备注：部署时间依据' in para.text:
        print(f"  段落{i}: {para.text[:80]}...")

print("\n=== 验证4.1理念创新 ===")
in_section = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '4.1 理念创新' in text:
        in_section = True
    if in_section and text:
        print(f"  [{i}] {text[:100]}")
    if in_section and '4.2 理论方法创新' in text:
        break
