#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""验证图3-1引用位置并修复"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-用于合并_修改版.docx')

print("=== 检查图3-1引用位置 ===")
for i, para in enumerate(doc.paragraphs):
    if '图3-1' in para.text or '图3-2' in para.text or '图3-3' in para.text:
        if not re.match(r'^图\s+3', para.text.strip()):
            print(f"[{i}] {para.text[:100]}")

# 检查3.6.1节内容
print("\n=== 3.6.1节内容 ===")
in_361 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.6.1' in text:
        in_361 = True
    if in_361 and text:
        print(f"[{i}] {text[:80]}")
    if in_361 and '3.6.2' in text:
        break
