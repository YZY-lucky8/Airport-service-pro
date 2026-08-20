#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""验证最终修改结果"""
from docx import Document
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_最终版.docx')

print("=== 验证问题1-5修正 ===")

# 问题1：布隆过滤器
text132 = doc.paragraphs[132].text
print(f"1. 段落132布隆过滤器:")
print(f"   位数组143773位: {'✅' if '143773' in text132 else '❌'}")
print(f"   10个哈希函数: {'✅' if '10个哈希' in text132 else '❌'}")
print(f"   双重哈希: {'✅' if '双重哈希' in text132 else '❌'}")
print(f"   18KB/万条: {'✅' if '18KB' in text132 else '❌'}")
print(f"   无MurmurHash3: {'✅' if 'MurmurHash3' not in text132 else '❌'}")

# 问题2：第五章
text526 = doc.paragraphs[526].text
print(f"\n2. 段落526第五章:")
print(f"   36项全部通过: {'✅' if '36项全部通过' in text526 else '❌'}")
print(f"   无33项26项: {'✅' if '33项' not in text526 else '❌'}")

# 问题3：内存
text114 = doc.paragraphs[114].text
print(f"\n3. 段落114内存:")
print(f"   77兆字节: {'✅' if '77兆' in text114 else '❌'}")
print(f"   无70兆: {'✅' if '70兆' not in text114 else '❌'}")

# 问题4：表22
cell22 = doc.tables[22].rows[3].cells[1].text
print(f"\n4. 表22布隆过滤器内存:")
print(f"   ≈18KB/万条: {'✅' if '18KB' in cell22 else '❌'}")
print(f"   无12KB: {'✅' if '12KB' not in cell22 else '❌'}")

# 问题5：段落206
text206 = doc.paragraphs[206].text
print(f"\n5. 段落206阈值配置:")
print(f"   预期容量10000条: {'✅' if '预期容量' in text206 else '❌'}")
print(f"   无10000位: {'✅' if '10000位' not in text206 else '❌'}")

print("\n=== 验证17篇参考文献角标 ===")
# 收集正文中所有角标
all_citations = set()
for i, para in enumerate(doc.paragraphs):
    # 跳过参考文献章节（533以后）
    if i >= 530:
        break
    matches = re.findall(r'\[(\d+)\]', para.text)
    for m in matches:
        all_citations.add(int(m))

print(f"正文中出现的角标: {sorted(all_citations)}")
missing = [i for i in range(1, 18) if i not in all_citations]
if missing:
    print(f"❌ 缺失的角标: {missing}")
else:
    print("✅ 17篇参考文献全部在正文中引用")

# 显示每个角标的位置
print("\n=== 角标位置详情 ===")
for i, para in enumerate(doc.paragraphs):
    if i >= 530:
        break
    matches = re.findall(r'\[(\d+)\]', para.text)
    if matches:
        print(f"  段落{i}: 引用{matches} - {para.text[:80]}...")
