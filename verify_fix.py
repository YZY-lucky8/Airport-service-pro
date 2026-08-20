#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""验证表格结果分析和页码设置"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告-(3)_修改版.docx')

print("=== 检查各测试表格后的内容 ===")
# 表题段落索引
table_captions = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^表\s+3-\d+', text):
        table_captions.append((i, text))

for idx, (para_idx, caption) in enumerate(table_captions):
    # 检查表题后1-3段
    next_texts = []
    for j in range(1, 4):
        if para_idx + j < len(doc.paragraphs):
            t = doc.paragraphs[para_idx + j].text.strip()
            if t:
                next_texts.append(t[:50])
    has_analysis = any('由表' in t or '测试' in t or '结果' in t or '通过' in t for t in next_texts)
    status = "✅" if has_analysis else "⚠️"
    print(f"  {status} {caption[:30]} -> 后段: {next_texts[:2]}")

print("\n=== 检查页码设置 ===")
for section in doc.sections:
    print(f"  节起始类型: {section.start_type}")
    print(f"  页边距: 上{section.top_margin}, 下{section.bottom_margin}")
    # 检查页码
    footer = section.footer
    if footer:
        for para in footer.paragraphs:
            if para.text.strip():
                print(f"  页脚内容: {para.text[:50]}")

print("\n=== 检查图2-9、图3-1引用 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '图2-9' in text:
        print(f"  图2-9引用: 段落{i} - {text[:80]}")
    if '图3-1' in text or '图3-2' in text or '图3-3' in text:
        print(f"  图3-x引用: 段落{i} - {text[:80]}")
