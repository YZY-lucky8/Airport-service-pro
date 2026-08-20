#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
doc = Document(r'C:\Users\Lenovo\Desktop\Airport-service-pro-master\作品报告_修改版_v2.docx')

print("=== 验证2.7.5节实测数据 ===")
in_section = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '2.7.5 性能指标真实性核验' in text:
        in_section = True
    if in_section and text:
        print(f"  [{i}] {text[:120]}")
    if in_section and ('2.8' in text or '第三章' in text):
        break

print("\n=== 验证1.2.2节简化表 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{ri}: {cells}")

print("\n=== 验证2.7.4节部署描述 ===")
for i, para in enumerate(doc.paragraphs):
    if '可直接部署于存量自助终端' in para.text and '2.7' not in para.text:
        print(f"  [{i}] {para.text[:150]}")
